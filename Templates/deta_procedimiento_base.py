"""
deta_procedimiento_base.py — v2.0
──────────────────────────────────────────────────────────────────────
Template base para procedimientos de cliente con identidad DETA.
Genera DOCX con python-docx y convierte a PDF vía docx2pdf o LibreOffice.

Uso:
    from deta_procedimiento_base import generar_procedimiento
    ruta_docx, ruta_pdf = generar_procedimiento(datos, output_dir)

Estructura del documento (9 secciones):
    1. Objetivo
    2. Alcance
    3. Políticas
    4. Definiciones
    5. Responsabilidades
    6. Propiedad
    7. Procedimiento  (subsecciones por tipo de caso)
    8. Historial de versiones  (tabla auto-generada)
    9. Flujo  (placeholder Lucidchart)

Identidad visual:
    - Títulos: Georgia, Navy #0c2b40
    - Body: Calibri 11pt, space_after 6pt
    - H2: 15pt Bold + gold underline + page_break_before en secciones 7-9
    - H3: 13pt Bold
    - H4: 11pt Bold Italic
    - Tabla historial/definiciones: header Gold con texto Navy
    - Header: logo DETA izquierda + nombre del procedimiento derecha
    - Footer: "DETA Consultores · detaconsultores.com · Confidencial" + paginación
    - Portada: bloque de identificación con título, cliente, versión, fecha

Cambios v2.0:
    - Logo path priorizado: Agent/Templates y Marca como primeros candidatos macOS
    - Bloque de portada al inicio del documento
    - Tamaños de heading explícitos (H2=15pt, H3=13pt, H4=11pt)
    - space_before/after controlado en todos los elementos
    - Page breaks automáticos antes de secciones 7, 8, 9
    - Gold underline en H2 headings
    - Anchos de columna en tablas (Término 4cm / Definición 12cm)
    - Filas alternadas en tablas de definiciones y responsabilidades
    - Header con fallback de texto si el logo no se encuentra

IMPORTANTE: No redefinir colores ni fuentes — usar las constantes de este módulo.

Dependencias:
    pip install python-docx docx2pdf
"""

import os
import subprocess
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── TOKENS DE COLOR ──────────────────────────────────────────────────────────

NAVY_HEX    = "0C2B40"
GOLD_HEX    = "D3AB6D"
SURFACE_HEX = "F5F7FA"
GRAY_HEX    = "6B7A8D"
WHITE_HEX   = "FFFFFF"

NAVY    = RGBColor(0x0C, 0x2B, 0x40)
GOLD    = RGBColor(0xD3, 0xAB, 0x6D)
SURFACE = RGBColor(0xF5, 0xF7, 0xFA)
GRAY    = RGBColor(0x6B, 0x7A, 0x8D)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

# Fila alternada en tablas
ROW_ALT_HEX = "F5F7FA"


# ─── RUTAS ────────────────────────────────────────────────────────────────────

def _resolve_logo(filename: str) -> str:
    """Detecta el entorno y retorna la ruta del logo. Prioriza macOS sobre Cowork."""
    candidates = [
        # macOS — Agent/Templates (tiene logos PNG de producción)
        os.path.join(
            "/Users/joelestrada/Library/CloudStorage/"
            "GoogleDrive-consultoria@detaconsultores.com/"
            "Mi unidad/Agent/Templates",
            filename,
        ),
        # macOS — Marca folder
        os.path.join(
            "/Users/joelestrada/Library/CloudStorage/"
            "GoogleDrive-consultoria@detaconsultores.com/"
            "Mi unidad/Agent/05_JOEL_OPERACION/Marca",
            filename,
        ),
        # macOS — .deta/assets (instalación local)
        os.path.expanduser(f"~/.deta/assets/{filename}"),
        # Cowork sandbox
        os.path.expanduser(f"~/mnt/Agent/Templates/skills/assets/{filename}"),
        f"/mnt/assets/{filename}",
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return candidates[2]  # fallback .deta/assets


LOGO_CYAN_PATH = _resolve_logo("logo-deta-cyan.png")


# ─── UTILIDADES XML ───────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Aplica color de fondo a una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.replace("#", ""))
    tcPr.append(shd)


def _add_tab_stop(para, pos_twips: int, alignment: str = "right"):
    """Agrega un tab stop a un párrafo."""
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), alignment)
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)


def _add_border_to_para(para, side: str, color_hex: str, sz: str = "6"):
    """Agrega borde a un lado de un párrafo (top/bottom/left/right)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), sz)
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), color_hex.replace("#", ""))
    pBdr.append(border)
    pPr.append(pBdr)


def _add_page_number_field(run):
    """Inserta campo PAGE en el run dado."""
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def _set_run_font(run, name: str, size_pt: float, color: RGBColor,
                  bold: bool = False, italic: bool = False):
    """Aplica fuente, tamaño, color y peso a un run."""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    # Forzar tamaño en XML para compatibilidad con estilos heredados
    rPr = run._r.get_or_add_rPr()
    for tag in ("w:sz", "w:szCs"):
        sz = OxmlElement(tag)
        sz.set(qn("w:val"), str(int(size_pt * 2)))  # half-points
        rPr.append(sz)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    rPr.append(color_el)


def _set_table_col_widths(table, widths_cm: list):
    """Asigna anchos de columna (en cm) a cada columna de la tabla."""
    for i, w in enumerate(widths_cm):
        for cell in table.column_cells(i):
            cell.width = Cm(w)


# ─── ESTRUCTURA DEL DOCUMENTO ─────────────────────────────────────────────────

def _setup_document() -> Document:
    """Crea un Document con márgenes DETA y estilos base."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # Estilo Normal base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = NAVY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return doc


def _setup_header(doc: Document, procedure_name: str):
    """Header: logo DETA izquierda, nombre del procedimiento derecha, línea gold."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    para = header.paragraphs[0]
    para.clear()

    # Logo o texto fallback izquierda
    logo_added = False
    if os.path.exists(LOGO_CYAN_PATH):
        try:
            run_logo = para.add_run()
            run_logo.add_picture(LOGO_CYAN_PATH, width=Inches(1.3))
            logo_added = True
        except Exception as e:
            import sys
            print(f"[deta_procedimiento_base] WARN logo: {e}", file=sys.stderr)

    if not logo_added:
        run_fallback = para.add_run("DETA Consultores")
        _set_run_font(run_fallback, "Georgia", 11, NAVY, bold=True)

    # Tab + nombre del procedimiento a la derecha
    run_name = para.add_run(f"\t{procedure_name}")
    _set_run_font(run_name, "Calibri", 9, GRAY)

    # Tab stop derecho ~16cm (9072 twips; 1cm ≈ 567 twips)
    _add_tab_stop(para, 9072, "right")

    # Línea gold inferior
    _add_border_to_para(para, "bottom", GOLD_HEX, sz="4")


def _setup_footer(doc: Document):
    """Footer: texto DETA izquierda, número de página derecha, línea gold superior."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    para = footer.paragraphs[0]
    para.clear()

    run_label = para.add_run("DETA Consultores · detaconsultores.com · Confidencial")
    _set_run_font(run_label, "Calibri", 8, GRAY)

    para.add_run("\t")

    run_pg = para.add_run()
    run_pg.font.name = "Calibri"
    run_pg.font.size = Pt(8)
    run_pg.font.color.rgb = GRAY
    _add_page_number_field(run_pg)

    _add_tab_stop(para, 9072, "right")
    _add_border_to_para(para, "top", GOLD_HEX, sz="4")


def _add_section_heading(doc: Document, text: str, level: int = 2,
                         page_break: bool = False):
    """
    Heading con identidad DETA.
    H2: Georgia 15pt Bold Navy + gold underline.
    H3: Georgia 13pt Bold Navy.
    H4: Calibri 11pt Bold Italic Navy.
    Acepta page_break=True para salto de página antes del heading.
    """
    para = doc.add_heading(text, level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Espaciado por nivel
    pf = para.paragraph_format
    if level == 2:
        pf.space_before = Pt(18)
        pf.space_after  = Pt(4)
    elif level == 3:
        pf.space_before = Pt(12)
        pf.space_after  = Pt(3)
    elif level == 4:
        pf.space_before = Pt(8)
        pf.space_after  = Pt(2)

    # Page break antes del párrafo
    if page_break:
        pPr = para._p.get_or_add_pPr()
        pb = OxmlElement("w:pageBreakBefore")
        pb.set(qn("w:val"), "1")
        pPr.append(pb)

    # Tamaño y estilo por nivel
    cfg = {
        1: (18.0, True,  False),
        2: (15.0, True,  False),
        3: (13.0, True,  False),
        4: (11.0, True,  True),
    }
    size, bold, italic = cfg.get(level, (11.0, False, False))
    font_name = "Georgia" if level <= 3 else "Calibri"

    for run in para.runs:
        _set_run_font(run, font_name, size, NAVY, bold=bold, italic=italic)

    # Gold underline para H2
    if level == 2:
        _add_border_to_para(para, "bottom", GOLD_HEX, sz="4")

    return para


def _add_body(doc: Document, text: str, italic: bool = False, shade: bool = False):
    """Párrafo de body Calibri 11pt Navy con spacing y justificación."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, "Calibri", 11, NAVY, italic=italic)
    para.paragraph_format.space_after  = Pt(6)
    para.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    if shade:
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), SURFACE_HEX)
        pPr.append(shd)
    return para


def _add_bullet(doc: Document, text: str):
    """Bullet point DETA."""
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    _set_run_font(run, "Calibri", 11, NAVY)
    para.paragraph_format.space_after  = Pt(3)
    para.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para


# ─── PORTADA / BLOQUE DE IDENTIFICACIÓN ───────────────────────────────────────

def _add_portada(doc: Document, datos: dict):
    """
    Bloque de identificación al inicio del documento.
    Eyebrow gold → Título grande → Línea gold → Meta (cliente/versión/fecha).
    """
    nombre   = datos.get("nombre", "Procedimiento")
    cliente  = datos.get("cliente", "")
    version  = datos.get("version", 1)

    # Eyebrow: "PROCEDIMIENTO" en gold uppercase
    para_ey = doc.add_paragraph()
    run_ey = para_ey.add_run("PROCEDIMIENTO")
    _set_run_font(run_ey, "Calibri", 8.5, GOLD, bold=True)
    para_ey.paragraph_format.space_before = Pt(0)
    para_ey.paragraph_format.space_after  = Pt(2)

    # Título principal en Georgia grande
    para_tit = doc.add_paragraph()
    run_tit = para_tit.add_run(nombre)
    _set_run_font(run_tit, "Georgia", 22, NAVY, bold=True)
    para_tit.paragraph_format.space_before = Pt(0)
    para_tit.paragraph_format.space_after  = Pt(6)
    _add_border_to_para(para_tit, "bottom", GOLD_HEX, sz="8")

    # Meta: cliente · versión · fecha
    fecha_str = datetime.now().strftime("%d-%b-%Y").replace(
        "Jan", "Ene").replace("Apr", "Abr").replace("Aug", "Ago").replace(
        "Dec", "Dic")
    meta = f"Cliente: {cliente}    ·    Versión: {version}    ·    Fecha: {fecha_str}"
    para_meta = doc.add_paragraph()
    run_meta = para_meta.add_run(meta)
    _set_run_font(run_meta, "Calibri", 9, GRAY)
    para_meta.paragraph_format.space_before = Pt(4)
    para_meta.paragraph_format.space_after  = Pt(20)


# ─── SECCIONES ────────────────────────────────────────────────────────────────

def _section_texto(doc: Document, num: int, title: str, content: str):
    """Sección simple de texto corrido."""
    _add_section_heading(doc, f"{num}. {title}")
    if content:
        _add_body(doc, content)
    else:
        _add_body(doc, "[Pendiente — completar con cliente]", italic=True)


def _section_lista(doc: Document, num: int, title: str, items: list):
    """Sección con lista de bullets."""
    _add_section_heading(doc, f"{num}. {title}")
    if items:
        for item in items:
            _add_bullet(doc, item)
    else:
        _add_body(doc, "[Pendiente — completar con cliente]", italic=True)


def _section_definiciones(doc: Document, definiciones: dict):
    """Sección 4: tabla de definiciones con anchos controlados y filas alternadas."""
    _add_section_heading(doc, "4. Definiciones")
    if not definiciones:
        _add_body(doc, "[Pendiente — completar con cliente]", italic=True)
        return

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _set_table_col_widths(table, [4.0, 12.0])

    # Header gold
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Término", "Definición"]):
        _set_cell_bg(cell, GOLD_HEX)
        run = cell.paragraphs[0].add_run(text)
        _set_run_font(run, "Calibri", 10, NAVY, bold=True)

    # Filas de datos con alternado
    for idx, (termino, definicion) in enumerate(definiciones.items()):
        row = table.add_row().cells
        bg = ROW_ALT_HEX if idx % 2 == 0 else WHITE_HEX
        for cell, text in zip(row, [termino, definicion]):
            _set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(text)
            _set_run_font(run, "Calibri", 10, NAVY)

    doc.add_paragraph()


def _section_responsabilidades(doc: Document, responsabilidades: dict):
    """Sección 5: tabla de responsabilidades con anchos controlados y filas alternadas."""
    _add_section_heading(doc, "5. Responsabilidades")
    if not responsabilidades:
        _add_body(doc, "[Pendiente — completar con cliente]", italic=True)
        return

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _set_table_col_widths(table, [4.0, 12.0])

    # Header gold
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Rol", "Responsabilidad"]):
        _set_cell_bg(cell, GOLD_HEX)
        run = cell.paragraphs[0].add_run(text)
        _set_run_font(run, "Calibri", 10, NAVY, bold=True)

    # Filas de datos con alternado
    for idx, (rol, desc) in enumerate(responsabilidades.items()):
        row = table.add_row().cells
        bg = ROW_ALT_HEX if idx % 2 == 0 else WHITE_HEX
        for cell, text in zip(row, [rol, desc]):
            _set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(text)
            _set_run_font(run, "Calibri", 10, NAVY)

    doc.add_paragraph()


def _section_procedimiento(doc: Document, secciones: list):
    """Sección 7: procedimiento con subsecciones y pasos. Page break al inicio."""
    _add_section_heading(doc, "7. Procedimiento", page_break=True)
    if not secciones:
        _add_body(doc, "[Pendiente — completar con cliente]", italic=True)
        return

    for i, seccion in enumerate(secciones, 1):
        titulo = seccion.get("titulo", f"Subsección {i}")
        pasos  = seccion.get("pasos", [])

        _add_section_heading(doc, f"7.{i} {titulo}", level=3)

        for paso in pasos:
            if isinstance(paso, dict):
                sub_titulo = paso.get("titulo", "")
                sub_pasos  = paso.get("pasos", [])
                _add_section_heading(doc, sub_titulo, level=4)
                for sp in sub_pasos:
                    _add_bullet(doc, sp)
            else:
                _add_bullet(doc, paso)


def _section_historial(doc: Document, version: int, cambios: str):
    """Sección 8: historial de versiones auto-generado. Page break al inicio."""
    _add_section_heading(doc, "8. Historial de versiones", page_break=True)

    headers = ["Versión", "Inicio vigencia", "Cambios sufridos",
               "Elaborado por:", "Revisado por:", "Autorizado por:"]

    table = doc.add_table(rows=2, cols=6)
    table.style = "Table Grid"
    _set_table_col_widths(table, [1.5, 2.0, 5.5, 3.0, 2.0, 2.0])

    # Header gold
    for cell, text in zip(table.rows[0].cells, headers):
        _set_cell_bg(cell, GOLD_HEX)
        run = cell.paragraphs[0].add_run(text)
        _set_run_font(run, "Calibri", 9, NAVY, bold=True)

    # Fila de datos
    fecha_hoy = datetime.now().strftime("%d-%b-%y").replace(
        "Jan", "Ene").replace("Apr", "Abr").replace("Aug", "Ago").replace(
        "Dec", "Dic")

    valores = [str(version), fecha_hoy, cambios or "Nueva creación",
               "DETA Consultores", "", ""]

    for cell, text in zip(table.rows[1].cells, valores):
        run = cell.paragraphs[0].add_run(text)
        _set_run_font(run, "Calibri", 9, NAVY)

    doc.add_paragraph()


def _section_flujo(doc: Document):
    """Sección 9: placeholder para diagrama Lucidchart. Page break al inicio."""
    _add_section_heading(doc, "9. Flujo", page_break=True)
    _add_body(
        doc,
        "[Diagrama de flujo pendiente — elaborar en Lucidchart e insertar aquí]",
        italic=True,
        shade=True,
    )


# ─── CONVERSIÓN A PDF ─────────────────────────────────────────────────────────

def _convert_to_pdf(docx_path: str) -> Optional[str]:
    """
    Convierte DOCX a PDF.
    Intenta docx2pdf primero (requiere MS Word en macOS/Windows).
    Fallback: LibreOffice headless.
    """
    pdf_path = docx_path.replace(".docx", ".pdf")
    output_dir = os.path.dirname(docx_path)

    # 1. docx2pdf (Word automation — mejor fidelidad)
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            return pdf_path
    except Exception as e:
        import sys
        print(f"[deta_procedimiento_base] docx2pdf falló: {e} — intentando LibreOffice",
              file=sys.stderr)

    # 2. LibreOffice headless (fallback)
    for soffice_bin in ["soffice", "/usr/bin/soffice",
                        "/Applications/LibreOffice.app/Contents/MacOS/soffice"]:
        try:
            subprocess.run(
                [soffice_bin, "--headless", "--convert-to", "pdf",
                 "--outdir", output_dir, docx_path],
                capture_output=True, text=True, timeout=60,
            )
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    import sys
    print("[deta_procedimiento_base] WARN: No se pudo generar PDF.", file=sys.stderr)
    return None


# ─── FUNCIÓN PRINCIPAL ────────────────────────────────────────────────────────

def generar_procedimiento(datos: dict, output_dir: str) -> tuple:
    """
    Genera el DOCX del procedimiento con identidad DETA y lo convierte a PDF.

    Args:
        datos: dict con claves:
            nombre, cliente, version, objetivo, alcance,
            politicas, definiciones, responsabilidades,
            propiedad, procedimiento, cambios_version
        output_dir: carpeta de destino

    Returns:
        tuple (ruta_docx, ruta_pdf) — ruta_pdf puede ser None si la conversión falló
    """
    os.makedirs(output_dir, exist_ok=True)

    nombre   = datos.get("nombre", "Procedimiento")
    cliente  = datos.get("cliente", "Cliente")
    version  = datos.get("version", 1)

    # Nomenclatura
    nombre_slug  = nombre.replace(" ", "_")
    cliente_slug = cliente.replace(" ", "_")
    filename     = f"Procedimiento_{nombre_slug}_{cliente_slug}_v{version}"
    docx_path    = os.path.join(output_dir, f"{filename}.docx")

    # Documento base
    doc = _setup_document()
    _setup_header(doc, nombre)
    _setup_footer(doc)

    # ── Portada / bloque de identificación ───────────────────────────────────
    _add_portada(doc, datos)

    # ── 1. Objetivo ──────────────────────────────────────────────────────────
    _section_texto(doc, 1, "Objetivo", datos.get("objetivo", ""))

    # ── 2. Alcance ───────────────────────────────────────────────────────────
    _section_texto(doc, 2, "Alcance", datos.get("alcance", ""))

    # ── 3. Políticas ─────────────────────────────────────────────────────────
    _section_lista(doc, 3, "Políticas", datos.get("politicas", []))

    # ── 4. Definiciones ──────────────────────────────────────────────────────
    _section_definiciones(doc, datos.get("definiciones", {}))

    # ── 5. Responsabilidades ─────────────────────────────────────────────────
    _section_responsabilidades(doc, datos.get("responsabilidades", {}))

    # ── 6. Propiedad ─────────────────────────────────────────────────────────
    _section_texto(doc, 6, "Propiedad", datos.get("propiedad", ""))

    # ── 7. Procedimiento ─────────────────────────────────── (page break) ───
    _section_procedimiento(doc, datos.get("procedimiento", []))

    # ── 8. Historial de versiones ──────────────────────────── (page break) ─
    _section_historial(doc, version, datos.get("cambios_version", ""))

    # ── 9. Flujo ────────────────────────────────────────────── (page break) ─
    _section_flujo(doc)

    # Guardar DOCX
    doc.save(docx_path)

    # Convertir a PDF
    pdf_path = _convert_to_pdf(docx_path)

    return docx_path, pdf_path


# ─── EJEMPLO DE USO ───────────────────────────────────────────────────────────

# if __name__ == "__main__":
#     datos_prueba = {
#         "nombre": "Salidas de Material",
#         "cliente": "SAPISA",
#         "version": 2,
#         "objetivo": "Estandarizar las salidas de material desde almacén.",
#         "alcance": "Aplica a todas las salidas de material originadas por el Procedimiento de Ventas.",
#         "politicas": [
#             "Ruta matutina: pedidos antes de las 17:30 del día anterior.",
#             "Evidencia obligatoria: todo documento debe firmarse y archivarse.",
#         ],
#         "definiciones": {
#             "OV": "Orden de Venta (documento que detona la salida).",
#             "Odoo": "Sistema para registrar OV, traslados, facturas y evidencias.",
#         },
#         "responsabilidades": {
#             "Almacén": "Valida la OV, revisa stock físico, prepara pedido.",
#             "Facturación": "Emite la factura cuando el flujo lo requiere.",
#         },
#         "propiedad": "Gerente de Almacén",
#         "procedimiento": [
#             {
#                 "titulo": "Inicio y Revisión de Stock",
#                 "pasos": [
#                     "Almacén → Revisa stock físico del material solicitado.",
#                     "Si stock incompleto → Notificar a ventas.",
#                     "Si stock completo → Definir tipo de movimiento.",
#                 ]
#             },
#         ],
#         "cambios_version": "Actualización según nuevo flujo de consigna.",
#     }
#     docx, pdf = generar_procedimiento(datos_prueba, "/tmp/")
#     print(f"DOCX: {docx}")
#     print(f"PDF:  {pdf}")
